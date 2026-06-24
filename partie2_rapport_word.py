"""
Projet Python Avancé – Partie 2
Analyse de "20 000 Lieues sous les mers" (Jules Verne) + génération rapport Word
"""

import requests
import re
import os
import math
import unittest
from collections import Counter
from io import BytesIO

import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches

from PIL import Image, ImageDraw, ImageFont, ImageOps
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ─── CONFIG ──────────────────────────────────────────────────────────────────

GUTENBERG_URL = "https://www.gutenberg.org/cache/epub/54873/pg54873.txt"
IMAGE_URL = "https://upload.wikimedia.org/wikipedia/commons/thumb/e/e1/Jules_Verne_-_Houghton_BMS_Fr_372.8_%281%29_-_cropped.jpg/640px-Jules_Verne_-_Houghton_BMS_Fr_372.8_%281%29_-_cropped.jpg"
# Fallback URLs si Wikipedia bloqué (None en dernier = génération auto)
IMAGE_FALLBACK_URLS = [
    "https://www.gutenberg.org/files/54873/54873-h/images/i_001.jpg",
    None,
]
LOGO_PATH = "logo_bw.png"
OUTPUT_DIR = "output_partie2"

os.makedirs(OUTPUT_DIR, exist_ok=True)


# ─── TÉLÉCHARGEMENT TEXTE ────────────────────────────────────────────────────

def download_book(url: str) -> str:
    """Télécharge le texte brut du livre depuis Gutenberg."""
    resp = requests.get(url, timeout=30)
    resp.raise_for_status()
    return resp.text


def extract_metadata(text: str) -> tuple[str, str]:
    """Extrait le titre et l'auteur depuis l'en-tête Gutenberg."""
    title, author = "20 000 Lieues sous les mers", "Jules Verne"
    for line in text.splitlines()[:40]:
        if line.startswith("Title:"):
            title = line.replace("Title:", "").strip()
        if line.startswith("Author:"):
            author = line.replace("Author:", "").strip()
    return title, author


def extract_first_chapter(text: str) -> str:
    """
    Extrait le premier chapitre du livre.
    Stratégie robuste : cherche le début du chapitre I puis le début du chapitre II.
    Fonctionne avec les différentes mises en page Gutenberg (FR/EN).
    """
    # Cherche le début du chapitre I (plusieurs variantes possibles)
    start_patterns = [
        r"CHAPITRE\s+(?:PREMIER|I)\b",
        r"CHAPTER\s+(?:FIRST|I)\b",
        r"Chapitre\s+(?:premier|I)\b",
    ]
    start_pos = -1
    for pat in start_patterns:
        m = re.search(pat, text, re.IGNORECASE)
        if m:
            start_pos = m.start()
            break

    if start_pos == -1:
        # Fallback : après l'en-tête Gutenberg
        marker = "*** START OF THE PROJECT GUTENBERG"
        idx = text.find(marker)
        start_pos = idx + 300 if idx != -1 else 500

    # Cherche le début du chapitre II pour délimiter
    end_patterns = [
        r"CHAPITRE\s+II\b",
        r"CHAPTER\s+II\b",
        r"Chapitre\s+II\b",
    ]
    end_pos = -1
    for pat in end_patterns:
        m = re.search(pat, text[start_pos + 100:], re.IGNORECASE)
        if m:
            end_pos = start_pos + 100 + m.start()
            break

    if end_pos == -1 or end_pos <= start_pos:
        end_pos = start_pos + 8000  # fallback : 8000 chars max

    chapter = text[start_pos:end_pos].strip()
    return chapter


# ─── ANALYSE PARAGRAPHES ─────────────────────────────────────────────────────

def split_paragraphs(chapter_text: str) -> list[str]:
    """
    Sépare le texte en paragraphes non vides.
    Normalise d'abord les fins de ligne Windows (\r\n) puis split sur lignes vides.
    """
    # Normalise TOUTES les variantes de fin de ligne → \n
    text = chapter_text.replace("\r\n", "\n").replace("\r", "\n")

    # Split sur une ou plusieurs lignes vides
    raw = re.split(r"\n{2,}", text)
    paragraphs = [p.strip().replace("\n", " ") for p in raw if len(p.strip()) > 20]
    return paragraphs


def count_words(paragraph: str) -> int:
    """Compte le nombre de mots dans un paragraphe."""
    return len(re.findall(r"\b\w+\b", paragraph))


def round_to_ten(n: int) -> int:
    """Arrondit à la dizaine inférieure (ex: 127 → 120)."""
    return (n // 10) * 10


def compute_paragraph_stats(paragraphs: list[str]) -> dict:
    """Calcule les statistiques demandées sur les paragraphes."""
    word_counts = [count_words(p) for p in paragraphs]
    rounded = [round_to_ten(w) for w in word_counts]
    rounded_sorted = sorted(rounded)
    distribution = Counter(rounded_sorted)

    total_words = sum(word_counts)
    return {
        "nb_paragraphs": len(paragraphs),
        "total_words": total_words,
        "min_words": min(word_counts) if word_counts else 0,
        "max_words": max(word_counts) if word_counts else 0,
        "avg_words": total_words / len(word_counts) if word_counts else 0,
        "word_counts": word_counts,
        "rounded_sorted": rounded_sorted,
        "distribution": distribution,
    }


# ─── GRAPHIQUE ───────────────────────────────────────────────────────────────

def make_chart(distribution: Counter, title: str, output_path: str) -> str:
    """Génère un graphique de distribution des longueurs de paragraphes."""
    labels = sorted(distribution.keys())
    values = [distribution[k] for k in labels]
    x_labels = [f"{l}" for l in labels]

    fig, ax = plt.subplots(figsize=(10, 5))
    fig.patch.set_facecolor("#1a1a2e")
    ax.set_facecolor("#16213e")

    bars = ax.bar(range(len(labels)), values, color="#e8002d", edgecolor="#ff6688", linewidth=0.8)
    ax.set_xticks(range(len(labels)))
    ax.set_xticklabels(x_labels, rotation=45, ha="right", fontsize=8, color="#e0e0e0")
    ax.set_xlabel("Nombre de mots (arrondi à la dizaine)", color="#e0e0e0", fontsize=11)
    ax.set_ylabel("Nombre de paragraphes", color="#e0e0e0", fontsize=11)
    ax.set_title(f"Distribution des longueurs de paragraphes\n{title}", color="white", fontsize=13, fontweight="bold")
    ax.tick_params(colors="#e0e0e0")
    for spine in ax.spines.values():
        spine.set_edgecolor("#333355")

    for bar, val in zip(bars, values):
        if val > 0:
            ax.text(bar.get_x() + bar.get_width() / 2, bar.get_height() + 0.1,
                    str(val), ha="center", va="bottom", color="white", fontsize=8)

    fig.tight_layout()
    fig.savefig(output_path, dpi=150, bbox_inches="tight", facecolor=fig.get_facecolor())
    plt.close(fig)
    return output_path


# ─── IMAGE ───────────────────────────────────────────────────────────────────

def _make_ocean_placeholder(output_path: str) -> str:
    """Crée une image d'illustration générée si aucune URL n'est accessible."""
    import math as _math
    img = Image.new("RGB", (800, 600), color=(10, 30, 80))
    draw = ImageDraw.Draw(img)
    for y in range(600):
        ratio = y / 600
        r = int(5 + 30 * ratio)
        g = int(20 + 90 * ratio)
        b = int(60 + 130 * ratio)
        draw.line([(0, y), (800, y)], fill=(r, g, b))
    for x in range(0, 800, 3):
        wy = 200 + int(25 * _math.sin(x * 0.025)) + int(12 * _math.sin(x * 0.06 + 1))
        draw.ellipse([x-2, wy-1, x+2, wy+1], fill=(180, 220, 255))
    draw.rectangle([0, 0, 800, 60], fill=(5, 15, 40))
    draw.text((400, 25), "20 000 Lieues sous les mers  –  Jules Verne", fill=(220, 200, 100), anchor="mm")
    draw.text((400, 560), "Illustration générée – Voyage sous les océans", fill=(160, 210, 255), anchor="mm")
    img.save(output_path, "JPEG", quality=85)
    return output_path


def download_and_process_image(url: str, output_path: str) -> str:
    """
    Télécharge l'image depuis url (avec fallbacks), recadre 4:3 et redimensionne.
    Si toutes les URLs échouent, génère une illustration de substitution.
    """
    urls_to_try = [url] + IMAGE_FALLBACK_URLS

    for try_url in urls_to_try:
        if try_url is None:
            print("  ℹ️  Génération d'une illustration de substitution…")
            return _make_ocean_placeholder(output_path)
        try:
            headers = {"User-Agent": "Mozilla/5.0 (student project; python-requests)"}
            resp = requests.get(try_url, timeout=20, headers=headers)
            resp.raise_for_status()
            img = Image.open(BytesIO(resp.content)).convert("RGB")

            # Recadrage centré 4:3
            w, h = img.size
            target_ratio = 4 / 3
            if w / h > target_ratio:
                new_w = int(h * target_ratio)
                left = (w - new_w) // 2
                img = img.crop((left, 0, left + new_w, h))
            else:
                new_h = int(w / target_ratio)
                top = (h - new_h) // 2
                img = img.crop((0, top, w, top + new_h))

            img = img.resize((800, 600), Image.LANCZOS)
            img.save(output_path, "JPEG", quality=85)
            print(f"  ✅ Image téléchargée depuis : {try_url}")
            return output_path
        except Exception as e:
            print(f"  ⚠️  URL échouée ({try_url[:60]}…) : {e}")
            continue

    return _make_ocean_placeholder(output_path)


def create_logo(output_path: str) -> str:
    """Crée un logo noir et blanc simple (si pas de logo existant)."""
    size = 200
    img = Image.new("L", (size, size), color=255)
    draw = ImageDraw.Draw(img)

    # Cercle extérieur
    draw.ellipse([10, 10, size - 10, size - 10], outline=0, width=6)
    # Texte
    draw.text((size // 2, size // 2 - 20), "F1", fill=0, anchor="mm")
    draw.text((size // 2, size // 2 + 20), "VERNE", fill=0, anchor="mm")
    img.save(output_path)
    return output_path


def compose_final_image(base_path: str, logo_path: str, output_path: str, angle: int = 25) -> str:
    """Colle le logo pivoté sur l'image principale."""
    base = Image.open(base_path).convert("RGBA")
    logo = Image.open(logo_path).convert("RGBA")

    # Redimensionner logo
    logo = logo.resize((120, 120), Image.LANCZOS)

    # Rotation du logo
    logo_rotated = logo.rotate(angle, expand=True)

    # Position : coin bas-droite avec marge
    bw, bh = base.size
    lw, lh = logo_rotated.size
    pos = (bw - lw - 20, bh - lh - 20)

    base.paste(logo_rotated, pos, logo_rotated)
    base.convert("RGB").save(output_path, "JPEG", quality=85)
    return output_path


# ─── RAPPORT WORD ─────────────────────────────────────────────────────────────

def _add_heading(doc: Document, text: str, level: int = 1):
    p = doc.add_heading(text, level=level)
    run = p.runs[0] if p.runs else p.add_run(text)
    run.bold = True
    if level == 1:
        run.font.size = Pt(20)
        run.font.color.rgb = RGBColor(0xE8, 0x00, 0x2D)
    elif level == 2:
        run.font.size = Pt(15)
        run.font.color.rgb = RGBColor(0x16, 0x21, 0x3E)
        run.italic = True
    return p


def generate_word_report(
    title: str,
    author: str,
    report_author: str,
    stats: dict,
    chart_path: str,
    final_image_path: str,
    output_path: str,
):
    """Génère le rapport Word complet."""
    doc = Document()

    # Marges
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.2)
        section.right_margin = Inches(1.2)

    # ── PAGE DE TITRE ────────────────────────────────────────────────────────
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("RAPPORT D'ANALYSE LITTÉRAIRE")
    run.bold = True
    run.font.size = Pt(24)
    run.font.color.rgb = RGBColor(0xE8, 0x00, 0x2D)

    doc.add_paragraph()  # Espace

    # Titre du livre
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    run.bold = True
    run.italic = True
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(0x16, 0x21, 0x3E)

    doc.add_paragraph()

    # Image principale
    if os.path.exists(final_image_path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(final_image_path, width=Inches(5.5))

    doc.add_paragraph()

    # Auteur du livre
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"Auteur : {author}")
    run.font.size = Pt(14)
    run.bold = True

    # Auteur du rapport
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"Rapport rédigé par : {report_author}")
    run.font.size = Pt(12)
    run.italic = True

    # Date
    from datetime import date
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"Date : {date.today().strftime('%d %B %Y')}")
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    doc.add_page_break()

    # ── PAGE GRAPHIQUE ───────────────────────────────────────────────────────
    _add_heading(doc, "Analyse du Premier Chapitre", level=1)

    _add_heading(doc, "Distribution des longueurs de paragraphes", level=2)

    if os.path.exists(chart_path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(chart_path, width=Inches(6))

    doc.add_paragraph()

    # ── DESCRIPTION ET STATISTIQUES ──────────────────────────────────────────
    _add_heading(doc, "Description et statistiques", level=2)

    desc = (
        f"L'œuvre analysée est \"{title}\" de {author}, roman d'aventures scientifiques "
        f"publié en 1870. Ce récit relate le voyage du professeur Aronnax à bord du "
        f"Nautilus, le sous-marin du mystérieux capitaine Nemo, à travers les profondeurs "
        f"des océans du monde entier."
    )
    p = doc.add_paragraph(desc)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    doc.add_paragraph()
    _add_heading(doc, "Données statistiques – Premier chapitre", level=2)

    table_data = [
        ("Nombre de paragraphes", str(stats["nb_paragraphs"])),
        ("Nombre total de mots", str(stats["total_words"])),
        ("Nombre minimal de mots (paragraphe)", str(stats["min_words"])),
        ("Nombre maximal de mots (paragraphe)", str(stats["max_words"])),
        ("Nombre moyen de mots (paragraphe)", f"{stats['avg_words']:.1f}"),
        ("Source des données", "Project Gutenberg (gutenberg.org)"),
        ("URL du texte", GUTENBERG_URL),
    ]

    table = doc.add_table(rows=len(table_data), cols=2)
    table.style = "Table Grid"

    for i, (label, value) in enumerate(table_data):
        row = table.rows[i]
        # Cellule label
        c1 = row.cells[0]
        p1 = c1.paragraphs[0]
        run1 = p1.add_run(label)
        run1.bold = True
        run1.font.size = Pt(10)
        # Cellule valeur
        c2 = row.cells[1]
        p2 = c2.paragraphs[0]
        run2 = p2.add_run(value)
        run2.font.size = Pt(10)
        # Couleur alternée
        if i % 2 == 0:
            for cell in (c1, c2):
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                shd = OxmlElement("w:shd")
                shd.set(qn("w:val"), "clear")
                shd.set(qn("w:color"), "auto")
                shd.set(qn("w:fill"), "EEF2FF")
                tcPr.append(shd)

    doc.save(output_path)
    print(f"✅ Rapport Word généré : {output_path}")
    return output_path


# ─── PIPELINE PRINCIPAL ───────────────────────────────────────────────────────

def run_partie2(report_author: str = "Mathis"):
    print("=" * 60)
    print("PARTIE 2 – Analyse littéraire + rapport Word")
    print("=" * 60)

    # 1. Télécharger le livre
    print("\n[1/7] Téléchargement du livre depuis Gutenberg…")
    try:
        text = download_book(GUTENBERG_URL)
    except Exception as e:
        print(f"  ❌ Erreur téléchargement : {e}")
        return

    # 2. Extraire métadonnées et chapitre
    print("[2/7] Extraction titre, auteur, premier chapitre…")
    title, author = extract_metadata(text)
    chapter = extract_first_chapter(text)
    print(f"  Titre  : {title}")
    print(f"  Auteur : {author}")
    print(f"  Chapitre extrait : {len(chapter)} caractères")
    print(f"  Aperçu début : {repr(chapter[:150])}")

    # 3. Analyse paragraphes
    print("[3/7] Analyse des paragraphes…")
    paragraphs = split_paragraphs(chapter)
    if not paragraphs:
        print("  ⚠️  Aucun paragraphe détecté — vérifiez l'extraction du chapitre")
        return
    stats = compute_paragraph_stats(paragraphs)
    print(f"  {stats['nb_paragraphs']} paragraphes | {stats['total_words']} mots")
    print(f"  Min: {stats['min_words']} | Max: {stats['max_words']} | Moy: {stats['avg_words']:.1f}")
    print(f"  Distribution (5 premiers) : {list(stats['distribution'].items())[:5]}")

    # 4. Graphique
    print("[4/7] Génération du graphique…")
    chart_path = os.path.join(OUTPUT_DIR, "chart_paragraphes.png")
    make_chart(stats["distribution"], title, chart_path)

    # 5. Télécharger et traiter image
    print("[5/7] Téléchargement et traitement de l'image…")
    img_raw_path = os.path.join(OUTPUT_DIR, "nautilus_raw.jpg")
    local_img = "JulesVerne.jpg"
    if os.path.exists(local_img):
        # Image locale disponible : on la recadre/redimensionne directement
        print("  ✅ Image locale trouvée : JulesVerne.jpg")
        from PIL import Image as _PIL
        img = _PIL.open(local_img).convert("RGB")
        w, h = img.size
        target_ratio = 4 / 3
        if w / h > target_ratio:
            new_w = int(h * target_ratio)
            left = (w - new_w) // 2
            img = img.crop((left, 0, left + new_w, h))
        else:
            new_h = int(w / target_ratio)
            top = (h - new_h) // 2
            img = img.crop((0, top, w, top + new_h))
        img = img.resize((800, 600), _PIL.LANCZOS)
        img.save(img_raw_path, "JPEG", quality=85)
    else:
        # Fallback : téléchargement depuis internet
        download_and_process_image(IMAGE_URL, img_raw_path)

    # 6. Logo + composition
    print("[6/7] Création logo et composition finale…")
    logo_path_local = os.path.join(OUTPUT_DIR, "logo_bw.png")
    if not os.path.exists(LOGO_PATH):
        create_logo(logo_path_local)
    else:
        logo_path_local = LOGO_PATH

    final_img_path = os.path.join(OUTPUT_DIR, "image_finale.jpg")
    try:
        compose_final_image(img_raw_path, logo_path_local, final_img_path, angle=30)
    except Exception as e:
        print(f"  ⚠️ Composition image : {e} — utilisation image brute")
        final_img_path = img_raw_path

    # 7. Rapport Word
    print("[7/7] Génération du rapport Word…")
    report_path = os.path.join(OUTPUT_DIR, "rapport_20000_lieues.docx")
    try:
        generate_word_report(
            title=title,
            author=author,
            report_author=report_author,
            stats=stats,
            chart_path=chart_path,
            final_image_path=final_img_path,
            output_path=report_path,
        )
    except Exception as e:
        print(f"  ❌ Erreur génération Word : {e}")
        return

    print("\n" + "=" * 60)
    print(f"✅ Pipeline terminé. Fichiers dans : {OUTPUT_DIR}/")
    print(f"   - {chart_path}")
    print(f"   - {final_img_path}")
    print(f"   - {report_path}")
    print("=" * 60)
    return report_path


# ─── TESTS UNITAIRES ─────────────────────────────────────────────────────────

class TestPartie2(unittest.TestCase):

    def test_count_words_basic(self):
        self.assertEqual(count_words("Bonjour le monde"), 3)

    def test_count_words_empty(self):
        self.assertEqual(count_words(""), 0)

    def test_count_words_punctuation(self):
        self.assertEqual(count_words("Bonjour, monde !"), 2)

    def test_round_to_ten(self):
        self.assertEqual(round_to_ten(123), 120)
        self.assertEqual(round_to_ten(127), 120)
        self.assertEqual(round_to_ten(129), 120)
        self.assertEqual(round_to_ten(130), 130)
        self.assertEqual(round_to_ten(100), 100)
        self.assertEqual(round_to_ten(9), 0)

    def test_split_paragraphs(self):
        text = "Premier paragraphe avec assez de mots.\n\nDeuxième paragraphe aussi long.\n\nCourt."
        paras = split_paragraphs(text)
        self.assertGreaterEqual(len(paras), 1)

    def test_compute_stats(self):
        paras = ["un deux trois", "quatre cinq six sept", "un deux"]
        stats = compute_paragraph_stats(paras)
        self.assertEqual(stats["nb_paragraphs"], 3)
        self.assertEqual(stats["min_words"], 2)
        self.assertEqual(stats["max_words"], 4)
        self.assertAlmostEqual(stats["avg_words"], 3.0, places=1)

    def test_compute_stats_empty(self):
        stats = compute_paragraph_stats([])
        self.assertEqual(stats["nb_paragraphs"], 0)
        self.assertEqual(stats["total_words"], 0)
        self.assertEqual(stats["avg_words"], 0)

    def test_extract_metadata_fallback(self):
        text = "Title: Mon Livre\nAuthor: Jean Dupont\n"
        title, author = extract_metadata(text)
        self.assertEqual(title, "Mon Livre")
        self.assertEqual(author, "Jean Dupont")

    def test_distribution_counter(self):
        paras = [
            "un " * 15,    # 15 → 10
            "un " * 25,    # 25 → 20
            "un " * 12,    # 12 → 10
        ]
        stats = compute_paragraph_stats(paras)
        dist = stats["distribution"]
        self.assertEqual(dist[10], 2)
        self.assertEqual(dist[20], 1)

    def test_chart_creates_file(self):
        dist = Counter({10: 3, 20: 1, 30: 2})
        path = "/tmp/test_chart.png"
        make_chart(dist, "Test", path)
        self.assertTrue(os.path.exists(path))
        os.remove(path)


def run_tests():
    loader = unittest.TestLoader()
    suite = loader.loadTestsFromTestCase(TestPartie2)
    runner = unittest.TextTestRunner(verbosity=2)
    result = runner.run(suite)
    return result.wasSuccessful()


# ─── MAIN ─────────────────────────────────────────────────────────────────────

if __name__ == "__main__":
    import sys
    if "--tests" in sys.argv:
        success = run_tests()
        sys.exit(0 if success else 1)
    else:
        run_partie2(report_author="Mathis")